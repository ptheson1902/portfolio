"""Export service for generating resume in PDF/DOCX formats."""
import os
from datetime import datetime
from typing import Literal, Dict, Any, List
from io import BytesIO
from sqlalchemy.orm import Session
from jinja2 import Environment, FileSystemLoader

from ..models.schemas import Language
from ..repositories.profile_repository import ProfileRepository
from ..repositories.skill_repository import SkillRepository
from ..repositories.project_repository import ProjectRepository


class ExportService:
    """Service for exporting resume to various formats."""

    def __init__(self, db: Session):
        self.db = db
        self.profile_repo = ProfileRepository(db)
        self.skill_repo = SkillRepository(db)
        self.project_repo = ProjectRepository(db)

        # Set up Jinja2 template environment
        templates_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
            "templates"
        )
        self.template_env = Environment(loader=FileSystemLoader(templates_dir))

    def _get_category_display_name(self, category_key: str, lang: Language) -> str:
        """Get display name for skill category."""
        display_names = {
            "programming_languages": {
                "ja": "プログラミング言語",
                "vi": "Ngôn ngữ lập trình",
                "en": "Programming Languages"
            },
            "frameworks": {
                "ja": "フレームワーク",
                "vi": "Framework",
                "en": "Frameworks"
            },
            "databases": {
                "ja": "データベース",
                "vi": "Cơ sở dữ liệu",
                "en": "Databases"
            },
            "cloud": {
                "ja": "クラウド",
                "vi": "Cloud",
                "en": "Cloud"
            },
            "ai_ml": {
                "ja": "AI/機械学習",
                "vi": "AI/Machine Learning",
                "en": "AI/Machine Learning"
            }
        }
        return display_names.get(category_key, {}).get(lang.value, category_key)

    def _gather_data(self, lang: Language) -> Dict[str, Any]:
        """Gather all data needed for the resume."""
        # Get profile data
        profile = self.profile_repo.get_profile()
        if not profile:
            raise ValueError("Profile not found")

        profile_data = {
            "name": profile.name,
            "name_kana": profile.name_kana,
            "name_vi": profile.name_vi,
            "age": profile.age,
            "school": profile.school,
            "graduation_year": profile.graduation_year,
            "field": profile.field,
            "work_experience": profile.work_experience,
            "japan_residence": profile.japan_residence,
            "japanese_level": profile.japanese_level,
            "self_pr": profile.self_pr.get(lang.value, profile.self_pr.get("ja", "")) if profile.self_pr else ""
        }

        # Get skills grouped by category with display names
        skills_raw = self.skill_repo.get_skills_grouped_by_category(lang, None)
        skills = {}
        for cat_key, skill_list in skills_raw.items():
            display_name = self._get_category_display_name(cat_key, lang)
            skills[display_name] = skill_list

        # Get projects
        projects = self.project_repo.get_projects_for_resume(lang)

        return {
            "profile": profile_data,
            "skills": skills,
            "projects": projects,
            "date": datetime.now().strftime("%Y-%m-%d")
        }

    def generate_html(self, lang: Language) -> str:
        """Generate HTML resume."""
        data = self._gather_data(lang)
        template = self.template_env.get_template(f"resume_{lang.value}.html")
        return template.render(**data)

    def generate_pdf(self, lang: Language) -> bytes:
        """Generate PDF resume using WeasyPrint."""
        try:
            from weasyprint import HTML
            html_content = self.generate_html(lang)
            return HTML(string=html_content).write_pdf()
        except ImportError:
            raise ImportError(
                "WeasyPrint is not installed. Install it with: pip install weasyprint. "
                "Note: WeasyPrint requires GTK libraries on Windows."
            )

    def generate_docx(self, lang: Language) -> bytes:
        """Generate DOCX resume using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")

        data = self._gather_data(lang)
        doc = Document()

        # Title
        title_labels = {"ja": "職務経歴書", "vi": "Sơ yếu lý lịch", "en": "Resume"}
        title = doc.add_heading(title_labels.get(lang.value, "Resume"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f"Date: {data['date']}")

        # Profile section
        summary_labels = {"ja": "職務要約", "vi": "Tóm tắt kinh nghiệm", "en": "Professional Summary"}
        doc.add_heading(summary_labels.get(lang.value, "Summary"), level=1)

        profile = data["profile"]
        profile_items = [
            ("Name", profile["name"]),
            ("Age", str(profile["age"])),
            ("Education", f"{profile['school']} ({profile['graduation_year']})"),
            ("Work Experience", profile["work_experience"]),
            ("Japan Residence", profile["japan_residence"]),
            ("Japanese Level", profile["japanese_level"]),
            ("Specialization", profile["field"])
        ]

        table = doc.add_table(rows=len(profile_items), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(profile_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        # Self PR
        about_labels = {"ja": "自己PR", "vi": "Giới thiệu bản thân", "en": "About Me"}
        doc.add_heading(about_labels.get(lang.value, "About"), level=1)
        doc.add_paragraph(profile["self_pr"])

        # Skills section
        skills_labels = {"ja": "スキル一覧", "vi": "Danh sách kỹ năng", "en": "Skills"}
        doc.add_heading(skills_labels.get(lang.value, "Skills"), level=1)

        for category, skills_list in data["skills"].items():
            doc.add_heading(category, level=2)
            skills_text = ", ".join([f"{s['name']} ({s['experience']})" for s in skills_list])
            doc.add_paragraph(skills_text)

        # Projects section
        exp_labels = {"ja": "業務経歴", "vi": "Kinh nghiệm làm việc", "en": "Work Experience"}
        doc.add_heading(exp_labels.get(lang.value, "Experience"), level=1)

        for project in data["projects"]:
            doc.add_heading(project["name"], level=2)
            doc.add_paragraph(
                f"Period: {project['start_date']} - {project.get('end_date', 'Present')} ({project['duration']})"
            )
            doc.add_paragraph(f"Role: {project['role']} | Team Size: {project['team_size']}")
            doc.add_paragraph(project["description"])

            doc.add_paragraph("Technologies: " + ", ".join(project["technologies"]))

            if project["highlights"]:
                for highlight in project["highlights"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(highlight)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_resume(
        self,
        lang: Language,
        format: Literal["pdf", "docx", "html"] = "pdf"
    ) -> bytes:
        """Generate resume in specified format."""
        if format == "pdf":
            return self.generate_pdf(lang)
        elif format == "docx":
            return self.generate_docx(lang)
        elif format == "html":
            return self.generate_html(lang).encode("utf-8")
        else:
            raise ValueError(f"Unsupported format: {format}")
